import re

from docx import Document
from docx.oxml import OxmlElement, ns
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.shared import RGBColor
from docx.shared import Cm


class MainRequirementsFormatter:

    @staticmethod
    def number_pages(doc):
        for section in doc.sections:
            section.footer.is_linked_to_previous = True

        new_paragraph = doc.sections[0].footer.add_paragraph()
        new_run = new_paragraph.add_run()

        new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.sections[0].different_first_page_header_footer = True

        fld_char1 = OxmlElement('w:fldChar')
        fld_char1.set(ns.qn('w:fldCharType'), 'begin')

        instr_text = OxmlElement('w:instrText')
        instr_text.set(ns.qn('xml:space'), 'preserve')
        instr_text.text = "PAGE"

        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(ns.qn('w:fldCharType'), 'end')

        new_run._r.append(fld_char1)
        new_run._r.append(instr_text)
        new_run._r.append(fld_char2)

    @staticmethod
    def change_font(font_name: str, paragraph):
        count = 0
        for run in paragraph.runs:
            if not run.font.name == font_name:
                run.font.name = font_name
                count += 1
        if (count > 0):
            return True
        else:
            return False

    @staticmethod
    def change_font_size(font_size: int, paragraph):
        count = 0
        for run in paragraph.runs:
            if not run.font.size == Pt(font_size):
                run.font.size = Pt(font_size)
                count += 1
        if (count > 0):
            return True
        else:
            return False


    @staticmethod
    def change_font_color(color: RGBColor, paragraph):
        count = 0
        for run in paragraph.runs:
            if not (run.font.color.rgb == color or run.font.color.rgb == None):
                run.font.color.rgb = color
                count+=1
        if (count>0):
            return True
        else:
            return False

    @staticmethod
    def change_line_spacing(spacing: WD_LINE_SPACING, paragraph):
        if not paragraph.paragraph_format.line_spacing_rule == spacing :
            paragraph.paragraph_format.line_spacing_rule = spacing
            return True
        else:
            return False

    @staticmethod
    def change_margins(left_cm, right_cm, top_cm, bottom_cm, doc: Document):
        count = 0
        for section in doc.sections:
            if not (section.left_margin //1000 == Cm(left_cm)//1000 and
                    section.right_margin//100000 == Cm(right_cm)//100000 and
                    section.top_margin//100 == Cm(top_cm)//100 and
                    section.bottom_margin//100 == Cm(bottom_cm)//100):
                section.left_margin = Cm(left_cm)
                section.right_margin = Cm(right_cm)
                section.top_margin = Cm(top_cm)
                section.bottom_margin = Cm(bottom_cm)
                count += 1
        if (count > 0):
            return True
        else:
            return False

    @staticmethod
    def change_left_paragraph_indentation(indentation_cm, paragraph):
        if not (paragraph.paragraph_format.left_indent == None):
            if not (round(paragraph.paragraph_format.left_indent.cm, 2) == indentation_cm):
                paragraph.paragraph_format.left_indent = Cm(indentation_cm)
                return True
        else:
            return False


    @staticmethod
    def change_title_page_year(doc: Document, year: str, changes):
        regex = re.compile(r"20[0-9][0-9]")

        for p in doc.paragraphs:
            for r in p.runs:
                if regex.search(r.text):
                    changes.append("year changed")
                    r.text = f'\n{year}'
                    return

    @staticmethod
    def external_changes(colour, indentation, font, size, spacing, margins, loc_changes):
        if colour:
            loc_changes.append("color changed")
        if indentation:
            loc_changes.append("indentation changed")
        if font:
            loc_changes.append("font name changed")
        if size:
            loc_changes.append("font size changed")
        if spacing:
            loc_changes.append("line spacing changed")
        if margins:
            loc_changes.append("margins changed")


    @staticmethod
    def format_document(doc: Document, loc_change):
        count_color_changes = False
        count_font_name_changes = False
        count_font_size_changes = False
        count_line_spacing_changes = False
        count_left_paragraph_indentation = False
        for paragraph in doc.paragraphs:
            count_font_name_changes = MainRequirementsFormatter.change_font('Times New Roman', paragraph)
            count_font_size_changes = MainRequirementsFormatter.change_font_size(14, paragraph)
            count_color_changes = MainRequirementsFormatter.change_font_color(RGBColor(0, 0, 0), paragraph)
            count_line_spacing_changes = MainRequirementsFormatter.change_line_spacing(WD_LINE_SPACING.ONE_POINT_FIVE, paragraph)
            count_left_paragraph_indentation = MainRequirementsFormatter.change_left_paragraph_indentation(1.25, paragraph)
        MainRequirementsFormatter.number_pages(doc)
        count_margins_changes = MainRequirementsFormatter.change_margins(3, 1.5, 2,2,doc)
        MainRequirementsFormatter.external_changes(count_color_changes, count_left_paragraph_indentation,count_font_name_changes,count_font_size_changes,count_line_spacing_changes,count_margins_changes, loc_change)
