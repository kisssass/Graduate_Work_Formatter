import re

from docx import Document


class SourceLinksFormatter:

    @staticmethod
    def get_number_of_ref_subscripts(doc: Document) -> int:
        counter = 0
        for p in doc.paragraphs:
            for r in p.runs:
                if r.font.superscript and re.match(r"\[[0-9]+]", r.text):
                    counter += 1
        return counter

    @staticmethod
    def get_numbered_list_id(paragraph):
        return paragraph._element.xpath('./w:pPr/w:numPr/w:numId/@w:val')

    @staticmethod
    def get_number_of_source_links(doc: Document, section_header: str) -> int:
        section_found = False
        first_list_id = None
        counter = 0
        for p in doc.paragraphs:
            if p.text.casefold() == section_header.casefold():
                section_found = True
            if section_found:
                list_id = SourceLinksFormatter.get_numbered_list_id(p)
                if list_id and not first_list_id:
                    first_list_id = list_id
                if first_list_id and list_id == first_list_id:
                    counter += 1

        return counter

    def extract_bibliography(doc: Document):
        bibliography_found = False
        bibliography_text = []
        changes = False

        for paragraph in doc.paragraphs:
            if bibliography_found:
                if paragraph.text.strip() == "":
                    break
                bibliography_text.append(paragraph.text)
            if "Библиографический список:" in paragraph.text:
                bibliography_found = True
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text in bibliography_text:
                #берём paragraph.text, передаём в скрипт, возвращаем изменённое
                #добавить if paragraph.text != изменённое, то изменяем файл(строка ниже doc.par...) и меняем changes на True
                #в самом конце функции, если changes==True то возвращаем тру и коммент ниже
                print(paragraph.text)
                doc.paragraphs[i].text = "read"

    @staticmethod
    def check_for_links_presence(doc: Document, changes):
        ref_subscripts_number = SourceLinksFormatter.get_number_of_ref_subscripts(doc)
        source_links_number = SourceLinksFormatter.get_number_of_source_links(doc, 'Библиографический список:')
        if ref_subscripts_number != source_links_number:
            changes.append("warning: the number of highlighted references does not correspond to the number of literature sources, perhaps your list of sources is not a complete list!")
        #тут if (extract_bibliography) то changes.append("отредактированы ссылки в списке библиографии")
        #если чуть больше заморочиться, то можно и уточнять какая ссылка, но мне кажется и этого достаточно